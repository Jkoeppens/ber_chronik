"""
parse_document.py — DOCX in strukturierte Segmente zerlegen

Unterstützte Dokumenttypen (--doc-type):
  buchnotizen    (Default) — hierarchische Buchnotizen mit level/source/page
  presseartikel  — flache Chronik / Presseartikel ohne Hierarchie

buchnotizen-Modus:

  Drei Strukturebenen:
  level 1  Projektüberschrift: DOCX-Heading-1-Zeilen, die als Gliederungsüberschriften
           dienen (keine Buchquellen). Aktuell: "Notizen" und "Übertrag von Zeitschriften".
           Erkennbar durch explizite Whitelist (ORGANIZER_H1); rein heuristische Erkennung
           scheitert an gemischt- und mehrsprachigen Überschriften dieses Dokuments.

  level 2  Literaturliste: Normal-Absätze direkt unter einem Organizer-Heading-1.
           type = 'bibliography'.

  level 3  Buch-Abschnitte: Normal-Absätze unter Heading-1-Buchquellen oder Heading-2.
           type = 'bibliography'  wenn Bibliographie-Signale erkannt (Kapitelverzeichnis,
                                  Zitate mit Anführungszeichen, Bibliothekssignaturen, …)
           type = 'content'       sonst  ← relevant für Zeitanker-Erkennung

page-Extraktion: 1–3 Ziffern am Zeilenende, getrennt durch Leerzeichen oder Klammer-zu.
source:    zuletzt gesetzte Buchüberschrift (Heading-1-Buchquelle oder Heading-2).
           Bei level-2-Segmenten: der Organizer-Heading-1-Text.

Output: data/interim/generalized/segments.json
"""

import argparse
import json
import re
import sys
from pathlib import Path

import docx

ROOT   = Path(__file__).resolve().parent.parent.parent
INPUT  = ROOT / "data" / "raw" / "Damakus Notizen.docx"
OUTPUT = ROOT / "data" / "interim" / "generalized" / "segments.json"

# ── Seitenzahl-Extraktion ──────────────────────────────────────────────────────
# 1–3 Ziffern am Ende, getrennt durch Leerzeichen oder Klammer-zu
PAGE_NR = re.compile(r"[\s)]+(\d{1,3})$")

# ── Bibliographie-Signale ──────────────────────────────────────────────────────
# Reihenfolge und Breite sind bewusst; werden OHNE 'page is None'-Guard angewendet,
# damit auch Kapitelverzeichnis-Einträge mit Startseite korrekt klassifiziert werden.
BIBLIO_RE = [
    re.compile(r"^[A-Z][a-z].{0,40}\s:\s[A-Z]"),    # "Author : Title" (Kapitelverzeichnis)
    re.compile(r'^["\u201c\u201e\u201a„(]'),           # beginnt mit Anführungszeichen / Klammer
    re.compile(r"\d:\s*[A-Za-z]{2}\s+\d+"),            # Bibliothekssignatur 4: Re 4618
    re.compile(r"\(\d{4}\)"),                           # Jahr in Klammern (1984)
    re.compile(r"\b(fernleihe|gelesen|nicht gefunden)\b", re.I),
]

# ── Organizer-Headings ─────────────────────────────────────────────────────────
# Heading-1-Überschriften, die keine Buchquellen sind, sondern Gliederungs-
# abschnitte einleiten. Normale Zeilen darunter gehen in die Literaturliste (level 2).
# Rein heuristische Erkennung (Länge, Sprache) ist in diesem Dokument unzuverlässig.
ORGANIZER_H1: set[str] = {
    "Notizen",
    "Übertrag von Zeitschriften",
}


def extract_page(text: str) -> tuple[int | None, str]:
    m = PAGE_NR.search(text)
    if m:
        return int(m.group(1)), text[: m.start()].rstrip()
    return None, text


def is_bibliography(text: str) -> bool:
    return any(p.search(text) for p in BIBLIO_RE)


def parse(path: Path) -> list[dict]:
    doc = docx.Document(path)
    segments: list[dict] = []
    seg_id = 0

    # state: 'bibliographic' = unter Organizer-H1 (→ level 2)
    #        'book_source'   = unter Buchquelle-H1 oder H2 (→ level 3)
    state          = "bibliographic"
    current_source: str | None = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name

        # ── Heading 1 ─────────────────────────────────────────────────────────
        if "Heading 1" in style:
            if text in ORGANIZER_H1:
                seg_id += 1
                segments.append({
                    "segment_id": f"s{seg_id:04d}",
                    "level":      1,
                    "type":       "meta",
                    "source":     None,
                    "text":       text,
                    "page":       None,
                })
                state          = "bibliographic"
                current_source = text
            else:
                state          = "book_source"
                current_source = text
            continue

        # ── Heading 2 ─────────────────────────────────────────────────────────
        if "Heading 2" in style:
            state          = "book_source"
            current_source = text
            continue

        # ── Normal / Normal (Web) / List Paragraph / … ────────────────────────
        page, clean = extract_page(text)

        if state == "bibliographic":
            level = 2
            typ   = "bibliography"
        else:
            level = 3
            typ   = "bibliography" if is_bibliography(text) else "content"

        seg_id += 1
        segments.append({
            "segment_id": f"s{seg_id:04d}",
            "level":      level,
            "type":       typ,
            "source":     current_source,
            "text":       clean,
            "page":       page,
        })

    return segments


# Kurze reine Jahres-Überschriften: weniger als 10 Zeichen, nur Ziffern (inkl. Leerzeichen)
_YEAR_HEADING = re.compile(r"^\d[\d\s]{0,8}$")


def parse_presseartikel(path: Path) -> list[dict]:
    """Flacher Modus für Presseartikel / Chronik-Dokumente.

    Jeder nicht-leere Absatz → eigenes Segment.
    Keine Hierarchie, kein source, kein page.
    Reine Jahres-Überschriften (< 10 Zeichen, nur Ziffern) → type 'heading'.
    Alle anderen → type 'content'.
    """
    doc = docx.Document(path)
    segments: list[dict] = []
    seg_id = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        if len(text) < 10 and _YEAR_HEADING.match(text):
            typ = "heading"
        else:
            typ = "content"

        seg_id += 1
        segments.append({
            "segment_id": f"s{seg_id:04d}",
            "type":       typ,
            "text":       text,
            "source":     None,
            "page":       None,
        })

    return segments


def main() -> None:
    from datetime import datetime
    ap = argparse.ArgumentParser(description="DOCX → segments.json")
    ap.add_argument("--project",  required=True,
                    help="Projektname (z.B. ber, damaskus)")
    ap.add_argument("--document", required=True,
                    help="Dokument-ID (z.B. main, uuid…)")
    ap.add_argument("input",      nargs="?", default=None,
                    help="Pfad zur DOCX-Datei (default: aus doc config.json)")
    ap.add_argument("--doc-type", default=None,
                    choices=["buchnotizen", "presseartikel"],
                    help="Dokumenttyp (default: buchnotizen)")
    args = ap.parse_args()

    project_dir = ROOT / "data" / "projects" / args.project
    doc_dir = project_dir / "documents" / args.document
    doc_dir.mkdir(parents=True, exist_ok=True)

    # Dokumentebene config lesen (falls schon vorhanden)
    doc_config: dict = {}
    doc_config_path = doc_dir / "config.json"
    if doc_config_path.exists():
        doc_config = json.loads(doc_config_path.read_text(encoding="utf-8"))

    # input_path: explizit > doc config > Fehler
    if args.input:
        input_path = Path(args.input)
        original_filename = Path(args.input).name
    elif doc_config.get("original_filename"):
        input_path = ROOT / "data" / "raw" / doc_config["original_filename"]
        original_filename = doc_config["original_filename"]
    else:
        print("Kein Input-Pfad angegeben und doc config.json hat kein original_filename.", file=sys.stderr)
        sys.exit(1)

    doc_type = args.doc_type or doc_config.get("doc_type", "buchnotizen")

    if not input_path.exists():
        print(f"Datei nicht gefunden: {input_path}", file=sys.stderr)
        sys.exit(1)

    if doc_type == "presseartikel":
        segments = parse_presseartikel(input_path)
    else:
        segments = parse(input_path)

    # doc_type als Metadatum in jedem Segment speichern
    for s in segments:
        s["doc_type"] = doc_type

    output_path = doc_dir / "segments.json"
    output_path.write_text(
        json.dumps(segments, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Dokument-config schreiben / aktualisieren
    doc_config.update({
        "doc_type":          doc_type,
        "original_filename": original_filename,
        "ingested_at":       datetime.now().isoformat(timespec="seconds"),
    })
    doc_config_path.write_text(
        json.dumps(doc_config, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    counts: dict[str, int] = {}
    for s in segments:
        counts[s["type"]] = counts.get(s["type"], 0) + 1

    print(f"{len(segments)} Segmente → {output_path}  (doc-type: {doc_type})")
    for typ, n in sorted(counts.items()):
        print(f"  {typ:12s} {n}")


if __name__ == "__main__":
    main()
